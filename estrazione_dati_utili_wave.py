

import os
import tempfile
import re
import openai
from docx import Document
from easyocr import Reader
from dotenv import load_dotenv

# Carica le variabili d'ambiente dal file .env
load_dotenv()

# Configurazione Azure/OpenAI
openai.api_type = os.getenv("OPENAI_API_TYPE")
openai.api_base = os.getenv("OPENAI_API_BASE")
openai.api_version = os.getenv("OPENAI_API_VERSION")
openai.api_key = os.getenv("OPENAI_API_KEY")
DEPLOYMENT_NAME = os.getenv("DEPLOYMENT_NAME")

# ============================================================================
# 1) Funzioni di normalizzazione
# ============================================================================

import hashlib

def normalize_text(text):
    """
    Rimuove spazi multipli, interruzioni extra e converte il testo in minuscolo.
    In questo modo, se l’input è logicamente uguale, la stringa normalizzata sarà identica.
    """
    return " ".join(text.split()).lower()

_cache = {}  # Cache globale, se vuoi puoi anche salvarla su file

def call_azure_openai_cached(full_text, system_prompt, user_prompt):
    """
    Esegue la chiamata ad Azure/OpenAI in maniera deterministica e utilizza una cache per
    restituire sempre lo stesso output se l'input e i prompt sono identici.
    """
    # Normalizza il testo in ingresso per creare una chiave coerente
    norm_text = normalize_text(full_text)
    key_input = norm_text + system_prompt + user_prompt
    key = hashlib.md5(key_input.encode("utf-8")).hexdigest()

    if key in _cache:
        print(f"[CACHE] Riutilizzo risultato per chiave: {key}")
        return _cache[key]
    else:
        # Esegui la chiamata (qui usiamo la funzione deterministica già esistente)
        result = call_azure_openai_deterministic(full_text, system_prompt, user_prompt)
        _cache[key] = result
        return result



# ============================================================================
# 1) Funzioni di estrazione testo e OCR
# ============================================================================
def extract_text_from_docx(docx_path):
    """
    Estrae testo da paragrafi e tabelle di un file DOCX.
    Ritorna una stringa unica con i contenuti (senza righe vuote).
    """
    doc = Document(docx_path)
    text_blocks = []

    # Legge i paragrafi
    for paragraph in doc.paragraphs:
        ptext = paragraph.text.strip()
        if ptext:
            text_blocks.append(ptext)

    # Legge le tabelle
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_blocks.append(" | ".join(row_text))

    return "\n".join(text_blocks)


def extract_images_from_docx(docx_path):
    """
    Estrae eventuali immagini dal DOCX, salvandole in una cartella temporanea.
    Ritorna la lista dei path alle immagini estratte.
    """
    image_paths = []
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            doc = Document(docx_path)
            for rel in doc.part.rels.values():
                try:
                    # Se la relazione è un'immagine
                    if "image" in rel.target_ref:
                        image_data = rel.target_part.blob
                        image_name = os.path.basename(rel.target_ref)
                        image_path = os.path.join(temp_dir, image_name)
                        with open(image_path, "wb") as f:
                            f.write(image_data)
                        image_paths.append(image_path)
                except Exception as e:
                    print(f"Immagine non valida ignorata: {e}")
        except Exception as e:
            print(f"Errore durante l'estrazione delle immagini: {e}")

        # Nota: al termine del blocco with, la temp dir viene rimossa,
        # per cui conviene se possibile processare i file subito
        # (vedi `ocr_on_images`).
    return image_paths


def ocr_on_images(image_paths):
    """
    Esegue OCR con EasyOCR su ciascuna immagine estratta e
    concatena il testo trovato in un'unica stringa.
    """
    reader = Reader(["it", "en"], gpu=False)
    extracted_texts = []

    for img_path in image_paths:
        try:
            results = reader.readtext(img_path, detail=0)  # detail=0 -> solo testo
            if results:
                extracted_texts.append("\n".join(results))
        except Exception as e:
            print(f"Errore OCR su immagine {img_path}: {e}")

    # Concatena testo OCR di tutte le immagini
    if extracted_texts:
        return "\n".join(extracted_texts)
    return ""


# ============================================================================
# 2) Funzione 'deterministica' per la chiamata ad Azure/OpenAI
# ============================================================================
def call_azure_openai_deterministic(full_text, system_prompt, user_prompt):
    """
    Esegue la richiesta al modello Chat (Azure/OpenAI) con parametri
    'deterministici' per ridurre la variabilità di output:
      - temperature=0.0
      - top_p=1.0
      - presence_penalty=0.0
      - frequency_penalty=0.0

    In caso di testi lunghi, suddivide in chunk e concatena i risultati
    per evitare di superare i limiti di token.
    """
    # Stima approssimata del numero di token
    approx_token_len = len(full_text) // 4  # ~ 1 token ogni 4 caratteri (euristica)
    TOKEN_LIMIT = 16000  # es. per GPT-4 o GPT-3.5 se supporta contesti ampi

    def single_chunk_call(txt_chunk):
        try:
            response = openai.ChatCompletion.create(
                engine=DEPLOYMENT_NAME,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt.format(content=txt_chunk)},
                ],
                temperature=0.0,
                top_p=1.0,
                presence_penalty=0.0,
                frequency_penalty=0.0,
                max_tokens=3000  # ipotesi -> puoi regolare
            )
            return response["choices"][0]["message"]["content"]
        except Exception as e:
            print(f"Errore nella chiamata OpenAI su chunk: {e}")
            return ""

    if approx_token_len <= TOKEN_LIMIT:
        # Se il testo è nei limiti, effettua un'unica chiamata
        return single_chunk_call(full_text).strip()
    else:
        # Suddivisione del testo in chunk
        chunk_size = TOKEN_LIMIT * 4  # *4 perché la stima token->caratteri
        start = 0
        partial_results = []

        while start < len(full_text):
            end = start + chunk_size
            chunk = full_text[start:end]
            partial_res = single_chunk_call(chunk)
            partial_results.append(partial_res.strip())
            start = end

        # Concatena tutti i parziali
        return "\n".join(partial_results).strip()


# ============================================================================
# 3) Funzione principale parse_aru_docx
# ============================================================================
def parse_aru_docx(docx_path):
    """
    1) Estrae testo dal DOCX (incluse immagini via OCR).
    2) Esegue:
       A) analisi interpretativa per i Function Point (IFPUG)
       B) un breve riassunto (~mezza pagina) su scopo dell'ARU
    3) Ritorna (fp_analysis, summary).
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Il file {docx_path} non esiste.")

    # Estrazione testo
    base_text = extract_text_from_docx(docx_path)
    image_paths = extract_images_from_docx(docx_path)

    # Esecuzione OCR se ci sono immagini
    if image_paths:
        ocr_text = ocr_on_images(image_paths)
        if ocr_text:
            base_text += "\n\n[TESTO ESTRATTO DA IMMAGINI]\n" + ocr_text

    # Prompt di sistema: contesto per l'analisi FP
    system_prompt_fp = (
        "Sei un analista esperto di Function Point Analysis (IFPUG). "
        "Il compito è interpretare il documento ARU, cercando: ILF, EIF, EI, EO, EQ, "
        "e le informazioni per stimare DET, RET e qualsiasi altro aspetto utile. "
        "Mantieni stabilità e chiarezza, senza introdurre dettagli casuali."
    )

    # Prompt utente: la richiesta di informazioni FP
    user_prompt_fp = (
        "Ecco il testo del documento (ARU):\n"
        "{content}\n\n"
        "1) Elenca le funzioni dati (ILF, EIF), se presenti. "
        "2) Elenca le funzioni transazionali (EI, EO, EQ) menzionate o inferibili. "
        "3) Fornisci indicazioni su DET/RET se possibile. "
        "4) Non unificare mai più sorgenti (EIF) se il documento le cita come separate. "
        "5) Se un requisito descrive più modalità di consultazione, classificale come EQ distinte. "
        "6) In generale, fornisci tutti i dettagli utili al calcolo dei function point, in modo coerente e ripetibile."
    )

    # Lato function points con caching per massimizzare la ripetibilità
    fp_analysis = call_azure_openai_cached(base_text, system_prompt_fp, user_prompt_fp)
    # Prompt di sistema: contesto per un riassunto
    system_prompt_summary = (
        "Sei un assistente che riassume il contenuto del documento. "
        "Non aggiungere nulla oltre a ciò che leggi."
    )

    # Prompt utente: generazione di mezza pagina di sintesi
    user_prompt_summary = (
        "Testo ARU:\n{content}\n\n"
        "Genera un riassunto di circa mezza pagina (max 200 parole) "
        "spiegando di cosa tratta la ARU, lo scopo del software e il contesto/committente."
    )

    half_page_summary = call_azure_openai_cached(base_text, system_prompt_summary, user_prompt_summary)
    _=[]
    return fp_analysis.strip(),_, half_page_summary.strip()


# ============================================================================
# 4) ESECUZIONE DI ESEMPIO
# ============================================================================
if __name__ == "__main__":
    # Sostituisci con il percorso reale del tuo file ARU .docx
    # docx_path = r"C:\Users\A395959\PycharmProjects\UFP_estimator\ARU_dir\ARU-Mercato-Re-factoringDamas(Analisi&DesignSprint17-18)_20240725103817.490_X.docx"
    docx_path = r"C:\Users\A395959\PycharmProjects\UFP_estimator\ARU_dir\SF_Piattaforma Unica VIS_v.1.0.docx"

    fp_info, _,summary = parse_aru_docx(docx_path)

    print("\n===== ANALISI FUNZIONI (FUNCTION POINT) =====")
    print(fp_info)

    print("\n===== SINTESI / RIASSUNTO ARU =====")
    print(summary)
