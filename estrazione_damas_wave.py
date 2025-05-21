import os
import re
import tempfile
import openai
import easyocr
from docx import Document
from dotenv import load_dotenv


# Carica variabili d'ambiente dal file .env (opzionale)
load_dotenv()

# Configura parametri OpenAI / Azure OpenAI
openai.api_type = os.getenv("OPENAI_API_TYPE")
openai.api_base = os.getenv("OPENAI_API_BASE")
openai.api_version = os.getenv("OPENAI_API_VERSION")
openai.api_key = os.getenv("OPENAI_API_KEY")
DEPLOYMENT_NAME = os.getenv("DEPLOYMENT_NAME")


# =========================================
# 1. Estrazione testo da DOCX (paragrafi + tabelle + immagini OCR)
# =========================================
def extract_text_from_docx(docx_path):
    """
    Estrae testo dai paragrafi e dalle tabelle di un file DOCX.
    Ritorna il testo concatenato.
    """
    try:
        doc = Document(docx_path)
        extracted_text = []

        # Estrarre testo dai paragrafi
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                extracted_text.append(text)

        # Estrarre testo dalle tabelle
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    extracted_text.append(" | ".join(row_text))

        return "\n".join(extracted_text)
    except Exception as e:
        print(f"Errore durante l'estrazione del testo da docx: {e}")
        return ""


def extract_images_from_docx(docx_path, temp_dir):
    """
    Estrae tutte le immagini dal file .docx, salvandole in temp_dir.
    Ritorna la lista dei percorsi delle immagini estratte.
    """
    try:
        doc = Document(docx_path)
        image_paths = []

        for rel in doc.part._rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_filename = os.path.join(temp_dir, os.path.basename(rel.target_ref))
                with open(image_filename, "wb") as img_file:
                    img_file.write(image_data)
                image_paths.append(image_filename)

        return image_paths
    except Exception as e:
        print(f"Errore durante l'estrazione delle immagini: {e}")
        return []


def extract_text_from_images(image_paths):
    """
    Esegue OCR sulle immagini in image_paths utilizzando EasyOCR.
    Ignora errori su immagini non leggibili.
    """
    try:
        reader = easyocr.Reader(['it', 'en'], gpu=False)
        ocr_texts = []
        for img_path in image_paths:
            try:
                result = reader.readtext(img_path, detail=0)
                ocr_texts.append(" ".join(result).strip())
            except Exception as e:
                print(f"Errore OCR su immagine {img_path}: {e}")
        return "\n".join(ocr_texts)
    except Exception as e:
        print(f"Errore OCR complessivo: {e}")
        return ""


def extract_all_content(docx_path):
    """
    Unisce testo da paragrafi/tabelle + testo estratto da immagini.
    """
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            text_paragraphs_tables = extract_text_from_docx(docx_path)
            image_paths = extract_images_from_docx(docx_path, temp_dir)
            text_images = extract_text_from_images(image_paths)

        full_text = text_paragraphs_tables
        if text_images:
            full_text += "\n\n[TESTO ESTRATTO DA IMMAGINI]\n" + text_images

        return full_text
    except Exception as e:
        print(f"Errore durante l'estrazione del contenuto: {e}")
        return ""


# =========================================
# 2. Rimozione eventuali indici / sommari (facoltativo)
# =========================================
def remove_index_from_text(full_text):
    """
    Rimuove parti di testo che assomigliano a un indice / sommario.
    """
    try:
        patterns = [
            re.compile(r"^\d+\.\s+.*\s+\d+$", re.MULTILINE),  # Esempio: "1. Introduzione  3"
            re.compile(r"^Indice.*$", re.IGNORECASE | re.MULTILINE),  # Linea che inizia con "Indice"
            re.compile(r"^Sommario.*$", re.IGNORECASE | re.MULTILINE),
            re.compile(r"^\d+\s+[A-Za-z].*\d+$", re.MULTILINE),  # "1 Titolo 1"
            re.compile(r"^[IVXLCDM]+\.\s+.*$", re.MULTILINE)  # Numerazione romana: "I. Titolo"
        ]

        filtered_text = full_text
        for pattern in patterns:
            filtered_text = re.sub(pattern, "", filtered_text)

        return filtered_text
    except Exception as e:
        print(f"Errore durante la rimozione dell'indice: {e}")
        return full_text


# =========================================
# 3. (Opzionale) Estrarre Requisiti Funzionali via Regex / Pattern
# =========================================
def extract_functional_requirements_regex(full_text):
    """
    Se il documento ha una struttura ricorrente (es: 'REQUISITI FUNZIONALI' ... 'FINE REQUISITI'),
    puoi usare un pattern per estrarre in modo deterministico.

    Ritorna la sottostringa corrispondente. Se non trovata, ritorna stringa vuota.

    Esempio di pattern, da adattare:
    """
    try:
        # Esempio (molto generico) di estrazione tra "Requisiti Funzionali" e "Fine Requisiti"
        pattern = re.compile(
            r"(REQUISITI FUNZIONALI[\s\S]*?)(FINE\s+REQUISITI|FINE REQ|\Z)",
            re.IGNORECASE
        )
        match = pattern.search(full_text)
        if match:
            # Prendiamo solo il primo gruppo
            return match.group(1).strip()
        else:
            return ""  # Nessun match
    except Exception as e:
        print(f"Errore durante l'estrazione via Regex: {e}")
        return ""


# =========================================
# 4. Estrarre Requisiti Funzionali via AI
#    (con Prompt deterministico e parametri fissi)
# =========================================
def extract_functional_requirements_with_ai(full_text):
    """
    Utilizza un modello AI per isolare i "Requisiti Funzionali".
    Prompt "ultra-deterministico" e parametri che azzerano la casualità.

    Nota: Se il documento è enorme e supera i limiti di token, si deve
    fare chunking con regole fisse e ripetibili.
    """
    try:
        # Prompt di sistema estremamente rigido
        system_message = (
            "Sei un assistente esperto nell'analisi dei documenti ARU. "
            "Devi restituire SOLO ed ESCLUSIVAMENTE il testo che corrisponde alla sezione "
            "\"Requisiti Funzionali\". Non inserire nulla di aggiuntivo, non fare riformulazioni "
            "o riassunti. Copia il testo letteralmente. Se non trovi nulla, restituisci stringa vuota."
        )

        # --------------------------------------------------------------------
        # 4a) Evita chunking se la lunghezza < soglia (e.g., 7000 token).
        #     In caso di modelli con contesti più ampi (es. GPT-3.5 16k),
        #     puoi regolare la soglia.
        #     NB: qui non facciamo conti di token precisi, è un esempio.
        # --------------------------------------------------------------------
        # Verifica lunghezza approssimativa in token (semplificato)
        approx_token_len = int(len(full_text) / 4)  # stima rozza (1 token ~ 4 caratteri)
        TOKEN_LIMIT = 7000

        if approx_token_len <= TOKEN_LIMIT:
            # Chiediamo direttamente con un singolo prompt
            response = openai.ChatCompletion.create(
                engine=DEPLOYMENT_NAME,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": full_text}
                ],
                max_tokens=3000,
                temperature=0.0,
                top_p=1.0,
                presence_penalty=0,
                frequency_penalty=0
            )
            extracted = response['choices'][0]['message']['content'].strip()
            return extracted
        else:
            # ----------------------------------------------------------------
            # 4b) Se il testo è troppo grande, dividiamo in chunk fissi.
            #     Mantieni SEMPRE la stessa dimensione, stesso ordine.
            # ----------------------------------------------------------------
            chunk_size = 6000  # Approssimazione in "caratteri"
            # Se vuoi più rigore, usa un contatore di token esatto (es: tiktoken).
            chunks = []
            start = 0
            while start < len(full_text):
                end = start + chunk_size
                chunk = full_text[start:end]
                chunks.append(chunk)
                start = end

            # Ora estraiamo i requisiti da ogni chunk e uniamo.
            # Attenzione: potresti ricevere "pezzi" di testo tronchi tra chunk.
            # Se la sezione Requisiti Funz. si spezza su più chunk, c’è da definire
            # come gestire la continuità. Per semplicità, concateno direttamente i risultati.
            extracted_sections = []
            for c in chunks:
                response = openai.ChatCompletion.create(
                    engine=DEPLOYMENT_NAME,
                    messages=[
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": c}
                    ],
                    max_tokens=2000,
                    temperature=0.0,
                    top_p=1.0,
                    presence_penalty=0,
                    frequency_penalty=0
                )
                part_extracted = response['choices'][0]['message']['content'].strip()
                extracted_sections.append(part_extracted)

            # Unisci in modo meccanico (fai attenzione a non perdere continuità se la sezione è tagliata a metà)
            # In modo super-semplice:
            return "\n".join(extracted_sections)

    except Exception as e:
        print(f"Errore durante l'estrazione AI dei requisiti funzionali: {e}")
        return "Errore AI durante l'estrazione."


# =========================================
# 5. Funzione Principale
# =========================================
def get_functional_requirements(docx_path, use_regex=False):
    """
    Data la path di un file docx, estrae il contenuto, filtra indici/sommari
    e cerca la sezione 'Requisiti Funzionali'.
    Se use_regex=True, usa un pattern fisso (se la struttura del doc è prevedibile).
    Altrimenti, usa AI in modo deterministico.
    """
    if not os.path.exists(docx_path):
        print(f"Errore: il file {docx_path} non esiste o non è accessibile.")
        return ""

    # 1) Estrai tutto
    full_content = extract_all_content(docx_path)
    # 2) Rimuovi eventuali indici / sommari
    filtered_content = remove_index_from_text(full_content)

    if use_regex:
        # Estraggo via Regex (se la struttura è nota)
        return extract_functional_requirements_regex(filtered_content)
    else:
        # Estraggo via AI con parametri deterministici
        return extract_functional_requirements_with_ai(filtered_content)


# =========================================
# 6. Esempio di esecuzione
# =========================================
if __name__ == "__main__":
    # docx_path = r"C:\Users\A395959\PycharmProjects\pyMilvus\ARU_dir\ARU-Mercato-Re-factoringDamas(Analisi&DesignSprint17-18)_20240725103817.490_X.docx"
    docx_path = r"C:\Users\A395959\PycharmProjects\UFP_estimator\ARU_dir\SF_Piattaforma Unica VIS_v.1.0.docx"


    # Se vuoi provare l'approccio AI:
    functional_requirements_ai = get_functional_requirements(docx_path, use_regex=False)
    print("=== REQUISITI FUNZIONALI (AI) ===")
    print(functional_requirements_ai)

    # Se vuoi provare l'approccio Regex (se la struttura del tuo doc lo consente):
    functional_requirements_regex = get_functional_requirements(docx_path, use_regex=True)
    print("=== REQUISITI FUNZIONALI (REGEX) ===")
    print(functional_requirements_regex)
